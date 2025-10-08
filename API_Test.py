from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Créer le document Word
doc = Document()

# Titre centré en gras
title = doc.add_heading("CONTRAT DE LOCATION D’UNE CHAMBRE", level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Entête
doc.add_paragraph("Entre les soussignés :\n")

doc.add_paragraph("Le bailleur :")
doc.add_paragraph("Nom : ___________________________")
doc.add_paragraph("Adresse : ________________________")
doc.add_paragraph("Téléphone : ______________________\n")

doc.add_paragraph("Et le locataire :")
doc.add_paragraph("Nom : ___________________________")
doc.add_paragraph("Adresse actuelle : ___________________")
doc.add_paragraph("Téléphone : ______________________\n")

# Article 1
doc.add_heading("Article 1 – Objet du contrat", level=2)
doc.add_paragraph("Le bailleur loue au locataire, qui accepte, une chambre située à l’adresse suivante :")
doc.add_paragraph("📍 _________________________________________________________________")
doc.add_paragraph("La chambre fait partie d’un logement (maison ou appartement) dont le bailleur est propriétaire ou locataire principal.")
doc.add_paragraph("Superficie de la chambre : ______ m²")
doc.add_paragraph("Elle comprend :")
doc.add_paragraph("☐ Un lit  ☐ Une armoire / placard  ☐ Un bureau  ☐ Une chaise")
doc.add_paragraph("☐ Autres : ____________________________________________")

# Article 2
doc.add_heading("Article 2 – Accès aux parties communes", level=2)
doc.add_paragraph("Le locataire aura libre accès aux pièces suivantes :")
doc.add_paragraph("☐ Cuisine  ☐ Salle de bain  ☐ Toilettes  ☐ Salon  ☐ Buanderie  ☐ Jardin")

# Article 3
doc.add_heading("Article 3 – Durée du contrat", level=2)
doc.add_paragraph("☐ Durée indéterminée, avec un préavis de ____ jours pour résiliation.")
doc.add_paragraph("☐ Durée déterminée : du ____/____/20__ au ____/____/20__.")

# Article 4
doc.add_heading("Article 4 – Loyer et charges", level=2)
doc.add_paragraph("💰 Loyer mensuel : ________ FCFA")
doc.add_paragraph("À payer d’avance, avant le __ de chaque mois.")
doc.add_paragraph("Charges (électricité, eau, internet, etc.) :")
doc.add_paragraph("☐ Incluses dans le loyer")
doc.add_paragraph("☐ À la charge du locataire")
doc.add_paragraph("☐ Forfait mensuel : ______ FCFA")
doc.add_paragraph("Mode de paiement :")
doc.add_paragraph("☐ Espèces  ☐ Mobile Money  ☐ Virement bancaire")

# Article 5
doc.add_heading("Article 5 – Dépôt de garantie", level=2)
doc.add_paragraph("Un dépôt de garantie de _______ FCFA est exigé à la signature du contrat.")
doc.add_paragraph("Il sera remboursé après état des lieux de sortie, déduction faite des éventuelles réparations.")

# Article 6
doc.add_heading("Article 6 – Obligations du locataire", level=2)
doc.add_paragraph("Le locataire s’engage à :")
doc.add_paragraph("- Payer le loyer à temps")
doc.add_paragraph("- Respecter le calme et la tranquillité")
doc.add_paragraph("- Ne pas sous-louer sans autorisation")
doc.add_paragraph("- Ne pas fumer dans les parties communes (sauf accord écrit)")
doc.add_paragraph("- Maintenir la chambre et les parties communes propres et en bon état")

# Article 7
doc.add_heading("Article 7 – Résiliation du contrat", level=2)
doc.add_paragraph("Chaque partie peut résilier le contrat avec un préavis de ____ jours.")
doc.add_paragraph("En cas de non-paiement ou de dommages graves, le bailleur peut résilier immédiatement.")

# Signatures
doc.add_paragraph("\nFait à ____________________, le ____/____/20__")
doc.add_paragraph("\nSignatures :\n")
doc.add_paragraph("Le Bailleur : ___________________________")
doc.add_paragraph("Le Locataire : ___________________________")

# Sauvegarde
file_path = "/C/users/Cheikh/Contrat_Location_Chambre.docx"
doc.save(file_path)
file_path